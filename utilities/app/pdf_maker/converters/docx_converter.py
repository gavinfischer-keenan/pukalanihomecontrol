"""
DOCX / DOC Converter
====================
Converts Microsoft Word documents to PDF using docx2pdf (requires MS Word on
Windows) with automatic detection and graceful fallback to text extraction via
python-docx if Word is not available.

Detection result is cached so the check only runs once per session.
"""

import os
import tempfile
import threading
from pathlib import Path
from typing import Optional

# Cache result of Word detection (None = not yet checked)
_word_status: Optional[tuple[bool, str]] = None
_detection_lock = threading.Lock()


def detect_word() -> tuple[bool, str]:
    """
    Check whether Microsoft Word (via docx2pdf) is available on this machine.
    Returns (available: bool, message: str).

    Result is cached after first call.
    """
    global _word_status
    with _detection_lock:
        if _word_status is not None:
            return _word_status

        # Try docx2pdf which wraps Word COM automation on Windows
        try:
            import docx2pdf  # noqa: F401

            # docx2pdf is installed — now confirm Word COM works
            try:
                import comtypes.client
                word_app = comtypes.client.CreateObject("Word.Application")
                word_app.Quit()
                _word_status = (True, "Microsoft Word detected and available.")
            except Exception:
                # comtypes failed (e.g., Word not installed)
                _word_status = (
                    False,
                    "docx2pdf is installed but Microsoft Word was not found on this system. "
                    "Word documents will be imported as plain text (formatting may differ).",
                )
        except ImportError:
            _word_status = (
                False,
                "docx2pdf is not installed. "
                "Word documents will be imported as plain text (formatting may differ).",
            )

        return _word_status


def convert_docx_to_pages(
    docx_path: str,
    page_size: str = "A4",
) -> tuple[list[str], list[str]]:
    """
    Convert a DOCX/DOC file to individual single-page temp PDFs.

    Strategy:
      1. Try docx2pdf (full fidelity, requires MS Word).
      2. Fall back to python-docx text extraction → ReportLab text PDF.
      3. If all else fails, return a single error page.

    Returns:
        (list_of_temp_pdf_paths, list_of_warning_strings)
    """
    warnings: list[str] = []
    word_available, word_msg = detect_word()

    if not word_available:
        warnings.append(f"⚠️ {word_msg}")

    # ------------------------------------------------------------------ #
    # Attempt 1: docx2pdf (Word COM automation)
    # ------------------------------------------------------------------ #
    if word_available:
        try:
            from docx2pdf import convert as docx2pdf_convert

            tmp_dir = tempfile.mkdtemp()
            out_pdf = os.path.join(tmp_dir, "converted.pdf")
            docx2pdf_convert(docx_path, out_pdf)

            if os.path.isfile(out_pdf) and os.path.getsize(out_pdf) > 0:
                from .text_converter import _split_pdf_to_pages
                pages = _split_pdf_to_pages(out_pdf, cleanup_source=True)
                return pages, warnings

            warnings.append("⚠️ docx2pdf produced an empty file — falling back to text extraction.")
        except Exception as e:
            warnings.append(f"⚠️ Word conversion error: {e}. Falling back to text extraction.")

    # ------------------------------------------------------------------ #
    # Attempt 2: python-docx text extraction → ReportLab
    # ------------------------------------------------------------------ #
    try:
        from docx import Document as DocxDocument  # python-docx

        doc = DocxDocument(docx_path)
        lines: list[str] = []

        for para in doc.paragraphs:
            lines.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells]
                lines.append("  |  ".join(t for t in row_texts if t))

        text_content = "\n".join(lines)

        # Write to temp text file and reuse text converter
        tmp_txt = tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False, encoding="utf-8"
        )
        tmp_txt.write(text_content or "[Document appears to be empty]")
        tmp_txt_path = tmp_txt.name
        tmp_txt.close()

        from .text_converter import convert_text_to_pages

        pages, txt_warnings = convert_text_to_pages(tmp_txt_path, page_size)
        warnings.extend(txt_warnings)

        try:
            os.unlink(tmp_txt_path)
        except Exception:
            pass

        if not word_available:
            warnings.append(
                "ℹ️ Rendered as plain text. Tables, images, and formatting inside "
                "the Word document may not appear correctly."
            )
        return pages, warnings

    except ImportError:
        warnings.append(
            "❌ python-docx is not installed. Cannot extract text from Word documents."
        )
    except Exception as e:
        warnings.append(f"❌ Text extraction failed: {e}")

    # ------------------------------------------------------------------ #
    # Attempt 3: Error page
    # ------------------------------------------------------------------ #
    from .pdf_handler import _make_error_page

    err_path = _make_error_page(
        f"Could not convert Word document:\n{Path(docx_path).name}\n\n"
        "Please ensure Microsoft Word or python-docx is installed."
    )
    return [err_path], warnings
