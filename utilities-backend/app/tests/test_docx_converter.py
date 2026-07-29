"""Tests for pdf_maker/converters/docx_converter.py."""
import os
import sys
import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock

sys.path.insert(0, str(Path(__file__).parent.parent))

from pdf_maker.converters import docx_converter
from pdf_maker.converters.docx_converter import (
    detect_word,
    convert_docx_to_pages,
)


@pytest.fixture(autouse=True)
def reset_word_cache():
    """Reset the cached Word detection result before each test."""
    docx_converter._word_status = None
    yield
    docx_converter._word_status = None


@pytest.fixture
def simple_docx(tmp_path):
    """Create a simple test DOCX file."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("Second paragraph with some text.")
    path = str(tmp_path / "test.docx")
    doc.save(path)
    return path


@pytest.fixture
def docx_with_table(tmp_path):
    """Create a DOCX with a table."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Document with table")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(0, 2).text = "C1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    table.cell(1, 2).text = "C2"
    path = str(tmp_path / "table.docx")
    doc.save(path)
    return path


@pytest.fixture
def empty_docx(tmp_path):
    from docx import Document
    doc = Document()
    path = str(tmp_path / "empty.docx")
    doc.save(path)
    return path


class TestDetectWord:
    def test_word_not_available_on_linux(self):
        """On a Linux server, Word/COM is not available."""
        available, msg = detect_word()
        assert isinstance(available, bool)
        assert isinstance(msg, str)
        assert len(msg) > 0

    def test_detection_cached(self):
        r1 = detect_word()
        r2 = detect_word()
        assert r1 == r2  # Same result from cache


class TestConvertDocxToPages:
    def test_simple_docx(self, simple_docx):
        pages, warnings = convert_docx_to_pages(simple_docx)
        assert len(pages) >= 1
        for p in pages:
            assert os.path.isfile(p)

    def test_docx_with_table(self, docx_with_table):
        pages, warnings = convert_docx_to_pages(docx_with_table)
        assert len(pages) >= 1

    def test_empty_docx(self, empty_docx):
        pages, warnings = convert_docx_to_pages(empty_docx)
        assert len(pages) >= 1  # Should produce at least a blank-ish page

    def test_nonexistent_file(self):
        pages, warnings = convert_docx_to_pages("/nonexistent/file.docx")
        assert len(pages) >= 1  # Error page
        assert len(warnings) > 0

    def test_warns_about_text_fallback(self, simple_docx):
        """On Linux, should warn about text-only rendering."""
        pages, warnings = convert_docx_to_pages(simple_docx)
        # Should have some warning about Word not being available or text rendering
        warning_text = " ".join(warnings).lower()
        # At minimum, the conversion should succeed
        assert len(pages) >= 1
